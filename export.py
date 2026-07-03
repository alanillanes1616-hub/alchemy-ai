"""
export.py — Exportación de conversaciones a TXT, Markdown, DOCX y PDF.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from memory import ConversationMemory
from utils import get_logger, timestamp_filename

logger = get_logger("export")


# ── TXT ───────────────────────────────────────────────────────────────────────

def to_txt(memory: ConversationMemory) -> tuple[bytes, str]:
    """Exporta el historial como texto plano."""
    content = memory.to_plain_text()
    filename = timestamp_filename("alchemy_chat", "txt")
    return content.encode("utf-8"), filename


# ── Markdown ──────────────────────────────────────────────────────────────────

def to_markdown(memory: ConversationMemory) -> tuple[bytes, str]:
    """Exporta el historial como Markdown."""
    content = memory.to_markdown()
    filename = timestamp_filename("alchemy_chat", "md")
    return content.encode("utf-8"), filename


# ── DOCX ──────────────────────────────────────────────────────────────────────

def to_docx(memory: ConversationMemory) -> tuple[bytes, str]:
    """Exporta el historial como documento Word."""
    doc = DocxDocument()

    # Título
    title = doc.add_heading("Conversación — Alchemy AI", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph("")

    for msg in memory.messages:
        label = "Usuario" if msg.role == "user" else "Alchemy AI"
        role_para = doc.add_paragraph()
        run = role_para.add_run(f"{label}  ·  {msg.timestamp}")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        content_para = doc.add_paragraph(msg.content)
        content_para.runs[0].font.size = Pt(11)
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    filename = timestamp_filename("alchemy_chat", "docx")
    return buffer.getvalue(), filename


# ── PDF ───────────────────────────────────────────────────────────────────────

def to_pdf(memory: ConversationMemory) -> tuple[bytes, str]:
    """Exporta el historial como PDF usando ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"],
        fontSize=16, textColor="#1D4ED8", spaceAfter=12,
    )
    label_style = ParagraphStyle(
        "Label", parent=styles["Normal"],
        fontSize=8, textColor="#6B7280", spaceBefore=8, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=6,
    )

    story = [
        Paragraph("Conversación — Alchemy AI", title_style),
        Spacer(1, 0.3 * cm),
    ]

    for msg in memory.messages:
        label = "Usuario" if msg.role == "user" else "Alchemy AI"
        story.append(Paragraph(f"{label}  ·  {msg.timestamp}", label_style))
        # Escapar caracteres especiales para ReportLab
        safe_content = (
            msg.content
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
        story.append(Paragraph(safe_content, body_style))
        story.append(Spacer(1, 0.2 * cm))

    doc.build(story)
    filename = timestamp_filename("alchemy_chat", "pdf")
    return buffer.getvalue(), filename


# ── Despachador público ───────────────────────────────────────────────────────

_EXPORTERS = {
    "TXT":      to_txt,
    "Markdown": to_markdown,
    "DOCX":     to_docx,
    "PDF":      to_pdf,
}

_MIME_TYPES = {
    "TXT":      "text/plain",
    "Markdown": "text/markdown",
    "DOCX":     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "PDF":      "application/pdf",
}


def export(memory: ConversationMemory, fmt: str) -> tuple[bytes, str, str]:
    """
    Exporta la conversación en el formato indicado.

    Args:
        memory: Instancia de ConversationMemory.
        fmt:    Uno de "TXT", "Markdown", "DOCX", "PDF".

    Returns:
        Tupla (bytes, filename, mime_type).

    Raises:
        ValueError: Si el formato no está soportado.
    """
    exporter = _EXPORTERS.get(fmt)
    if not exporter:
        raise ValueError(f"Formato de exportación no soportado: {fmt}")

    data, filename = exporter(memory)
    mime = _MIME_TYPES[fmt]
    logger.info("Exported conversation as %s (%d bytes).", fmt, len(data))
    return data, filename, mime
